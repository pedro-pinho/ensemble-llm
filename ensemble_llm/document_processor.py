"""Document Processing System for Ensemble LLM

Handles extraction, chunking, and processing of PDF, DOC, and DOCX documents.
Integrates with the memory system for long-term document storage and retrieval.
"""

import logging
import hashlib
from pathlib import Path
from typing import List, Dict, Any, Optional, Tuple
from datetime import datetime
import re
from dataclasses import dataclass

# PDF processing
try:
    import PyPDF2
    HAS_PYPDF2 = True
except ImportError:
    HAS_PYPDF2 = False

try:
    import pdfplumber
    HAS_PDFPLUMBER = True
except ImportError:
    HAS_PDFPLUMBER = False

# DOCX processing
try:
    from docx import Document
    HAS_PYTHON_DOCX = True
except ImportError:
    HAS_PYTHON_DOCX = False

# DOC processing (legacy)
try:
    import antiword
    HAS_ANTIWORD = True
except ImportError:
    HAS_ANTIWORD = False

logger = logging.getLogger("EnsembleLLM.DocumentProcessor")


@dataclass
class DocumentChunk:
    """A chunk of a document with metadata"""

    content: str
    chunk_index: int
    total_chunks: int
    metadata: Dict[str, Any]
    embedding_ready: bool = True

    def to_dict(self) -> Dict[str, Any]:
        return {
            "content": self.content,
            "chunk_index": self.chunk_index,
            "total_chunks": self.total_chunks,
            "metadata": self.metadata,
        }


@dataclass
class ProcessedDocument:
    """Represents a fully processed document"""

    document_id: str
    filename: str
    file_type: str
    total_pages: int
    total_chunks: int
    chunks: List[DocumentChunk]
    metadata: Dict[str, Any]
    processed_at: datetime

    def to_dict(self) -> Dict[str, Any]:
        return {
            "document_id": self.document_id,
            "filename": self.filename,
            "file_type": self.file_type,
            "total_pages": self.total_pages,
            "total_chunks": self.total_chunks,
            "metadata": self.metadata,
            "processed_at": self.processed_at.isoformat(),
        }


class DocumentProcessor:
    """Process documents for storage in memory system"""

    def __init__(
        self,
        chunk_size: int = 800,
        chunk_overlap: int = 200,
        min_chunk_size: int = 100,
    ):
        """
        Initialize document processor

        Args:
            chunk_size: Target size for each chunk (in tokens, ~4 chars/token)
            chunk_overlap: Number of tokens to overlap between chunks
            min_chunk_size: Minimum chunk size to avoid tiny fragments
        """
        self.chunk_size = chunk_size * 4  # Convert to characters
        self.chunk_overlap = chunk_overlap * 4
        self.min_chunk_size = min_chunk_size * 4

        # Check available libraries
        self._check_dependencies()

    def _check_dependencies(self):
        """Check which document processing libraries are available"""
        self.pdf_support = HAS_PYPDF2 or HAS_PDFPLUMBER
        self.docx_support = HAS_PYTHON_DOCX
        self.doc_support = HAS_ANTIWORD

        if not self.pdf_support:
            logger.warning(
                "No PDF processing library found. Install 'PyPDF2' or 'pdfplumber' "
                "for PDF support: pip install PyPDF2 pdfplumber"
            )

        if not self.docx_support:
            logger.warning(
                "python-docx not found. Install for DOCX support: "
                "pip install python-docx"
            )

        if not self.doc_support:
            logger.info(
                "antiword not found. Legacy DOC files may not be supported. "
                "For DOC support, install antiword or convert to DOCX."
            )

    def can_process(self, file_path: Path) -> bool:
        """Check if this file type can be processed"""
        suffix = file_path.suffix.lower()

        if suffix == ".pdf":
            return self.pdf_support
        elif suffix == ".docx":
            return self.docx_support
        elif suffix == ".doc":
            return self.doc_support

        return False

    def process_document(
        self,
        file_path: Path,
        metadata: Optional[Dict[str, Any]] = None,
    ) -> ProcessedDocument:
        """
        Process a document file into chunks

        Args:
            file_path: Path to the document file
            metadata: Additional metadata to attach to all chunks

        Returns:
            ProcessedDocument with all chunks and metadata

        Raises:
            ValueError: If file type not supported or processing fails
        """
        if not file_path.exists():
            raise ValueError(f"File not found: {file_path}")

        if not self.can_process(file_path):
            raise ValueError(
                f"Unsupported file type: {file_path.suffix}. "
                f"Supported: PDF={self.pdf_support}, DOCX={self.docx_support}, "
                f"DOC={self.doc_support}"
            )

        # Extract text and page info
        suffix = file_path.suffix.lower()

        if suffix == ".pdf":
            text, num_pages = self._extract_pdf(file_path)
        elif suffix == ".docx":
            text, num_pages = self._extract_docx(file_path)
        elif suffix == ".doc":
            text, num_pages = self._extract_doc(file_path)
        else:
            raise ValueError(f"Unsupported file type: {suffix}")

        # Generate document ID
        file_hash = hashlib.md5(file_path.read_bytes()).hexdigest()
        document_id = f"doc_{file_hash}_{datetime.now().timestamp()}"

        # Clean and chunk the text
        cleaned_text = self._clean_text(text)
        chunks = self._create_chunks(cleaned_text)

        # Create metadata for each chunk
        base_metadata = {
            "document_id": document_id,
            "filename": file_path.name,
            "file_type": suffix,
            "total_pages": num_pages,
            "upload_date": datetime.now().isoformat(),
        }

        # Merge with provided metadata
        if metadata:
            base_metadata.update(metadata)

        # Create DocumentChunk objects
        document_chunks = []
        for idx, chunk_text in enumerate(chunks):
            chunk_metadata = base_metadata.copy()
            chunk_metadata.update({
                "chunk_index": idx,
                "total_chunks": len(chunks),
            })

            document_chunks.append(
                DocumentChunk(
                    content=chunk_text,
                    chunk_index=idx,
                    total_chunks=len(chunks),
                    metadata=chunk_metadata,
                )
            )

        logger.info(
            f"Processed {file_path.name}: {num_pages} pages, "
            f"{len(chunks)} chunks, {len(cleaned_text)} chars"
        )

        return ProcessedDocument(
            document_id=document_id,
            filename=file_path.name,
            file_type=suffix,
            total_pages=num_pages,
            total_chunks=len(chunks),
            chunks=document_chunks,
            metadata=base_metadata,
            processed_at=datetime.now(),
        )

    def _extract_pdf(self, file_path: Path) -> Tuple[str, int]:
        """Extract text from PDF file"""
        text = ""
        num_pages = 0

        # Try pdfplumber first (better text extraction)
        if HAS_PDFPLUMBER:
            try:
                import pdfplumber
                with pdfplumber.open(file_path) as pdf:
                    num_pages = len(pdf.pages)
                    for page in pdf.pages:
                        page_text = page.extract_text()
                        if page_text:
                            text += page_text + "\n\n"

                if text.strip():
                    return text, num_pages
            except Exception as e:
                logger.warning(f"pdfplumber failed: {e}, trying PyPDF2")

        # Fallback to PyPDF2
        if HAS_PYPDF2:
            try:
                with open(file_path, 'rb') as file:
                    pdf_reader = PyPDF2.PdfReader(file)
                    num_pages = len(pdf_reader.pages)

                    for page in pdf_reader.pages:
                        page_text = page.extract_text()
                        if page_text:
                            text += page_text + "\n\n"

                return text, num_pages
            except Exception as e:
                raise ValueError(f"Failed to extract PDF: {e}")

        raise ValueError("No PDF processing library available")

    def _extract_docx(self, file_path: Path) -> Tuple[str, int]:
        """Extract text from DOCX file"""
        if not HAS_PYTHON_DOCX:
            raise ValueError("python-docx not installed")

        try:
            doc = Document(file_path)

            # Extract all paragraphs
            text = ""
            for para in doc.paragraphs:
                if para.text.strip():
                    text += para.text + "\n\n"

            # Extract text from tables
            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        if cell.text.strip():
                            text += cell.text + "\n"
                text += "\n"

            # Estimate pages (rough estimate: 500 words per page)
            word_count = len(text.split())
            num_pages = max(1, word_count // 500)

            return text, num_pages

        except Exception as e:
            raise ValueError(f"Failed to extract DOCX: {e}")

    def _extract_doc(self, file_path: Path) -> Tuple[str, int]:
        """Extract text from legacy DOC file"""
        # For legacy .doc files, we'd need antiword or similar
        # For simplicity, we'll suggest converting to DOCX
        raise ValueError(
            "Legacy .doc files are not supported. "
            "Please convert to .docx format using Microsoft Word or LibreOffice."
        )

    def _clean_text(self, text: str) -> str:
        """Clean and normalize extracted text"""
        # Remove excessive whitespace
        text = re.sub(r'\n{3,}', '\n\n', text)
        text = re.sub(r' {2,}', ' ', text)

        # Remove form feed and other control characters
        text = re.sub(r'[\x0c\x0b]', '', text)

        # Normalize line endings
        text = text.replace('\r\n', '\n')

        # Remove empty lines
        lines = [line.rstrip() for line in text.split('\n')]
        text = '\n'.join(lines)

        return text.strip()

    def _create_chunks(self, text: str) -> List[str]:
        """
        Split text into overlapping chunks

        Uses a sliding window approach with overlap to maintain context.
        """
        if len(text) <= self.chunk_size:
            return [text]

        chunks = []
        start = 0

        while start < len(text):
            # Get chunk
            end = start + self.chunk_size
            chunk = text[start:end]

            # If not at the end, try to break at sentence boundary
            if end < len(text):
                # Look for sentence endings within the last 20% of the chunk
                search_start = max(start + int(self.chunk_size * 0.8), start)
                search_text = text[search_start:end]

                # Find last sentence ending
                sentence_end = max(
                    search_text.rfind('. '),
                    search_text.rfind('! '),
                    search_text.rfind('? '),
                    search_text.rfind('.\n'),
                )

                if sentence_end != -1:
                    # Break at sentence boundary
                    end = search_start + sentence_end + 1
                    chunk = text[start:end]

            # Only add non-empty chunks of sufficient size
            if chunk.strip() and len(chunk.strip()) >= self.min_chunk_size:
                chunks.append(chunk.strip())

            # Move start position with overlap
            start = end - self.chunk_overlap

            # Prevent infinite loop
            if start <= chunks[-1] if chunks else 0:
                start = end

        return chunks

    def get_document_summary(self, processed_doc: ProcessedDocument) -> str:
        """Generate a summary of the document for display"""
        first_chunk = processed_doc.chunks[0].content if processed_doc.chunks else ""
        preview = first_chunk[:200] + "..." if len(first_chunk) > 200 else first_chunk

        return (
            f"Document: {processed_doc.filename}\n"
            f"Type: {processed_doc.file_type.upper()}\n"
            f"Pages: {processed_doc.total_pages}\n"
            f"Chunks: {processed_doc.total_chunks}\n"
            f"Preview: {preview}"
        )


# Convenience function
def process_document_file(
    file_path: Path,
    chunk_size: int = 800,
    chunk_overlap: int = 200,
    metadata: Optional[Dict[str, Any]] = None,
) -> ProcessedDocument:
    """
    Process a document file in one call

    Args:
        file_path: Path to document file
        chunk_size: Target chunk size in tokens
        chunk_overlap: Overlap between chunks in tokens
        metadata: Additional metadata to attach

    Returns:
        ProcessedDocument ready for memory storage
    """
    processor = DocumentProcessor(
        chunk_size=chunk_size,
        chunk_overlap=chunk_overlap,
    )

    return processor.process_document(file_path, metadata=metadata)
